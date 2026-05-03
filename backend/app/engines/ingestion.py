# C:\Users\grihi\OneDrive\Desktop\Personal Projects\Veridict\backend\app\engines\ingestion.py
from __future__ import annotations

import asyncio
import logging
import re
import tempfile
import uuid
from collections import defaultdict
from pathlib import Path

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from fastapi import UploadFile
from PIL import Image, ImageEnhance, ImageOps

from ..models import TextBlock

logger = logging.getLogger(__name__)

SUPPORTED_DOC_EXTS = {".pdf", ".docx"}
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}
PDF_SCANNED_MAX_CHARS = 49
MIN_BLOCK_CHARS = 10
OCR_CONTRAST_FACTOR = 2.0
BINARY_THRESHOLD = 128


def _upload_root() -> Path:
    return Path(tempfile.gettempdir()) / "veridict_uploads"


def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    contrasted = ImageEnhance.Contrast(gray).enhance(OCR_CONTRAST_FACTOR)
    return contrasted.point(lambda p: 255 if p >= BINARY_THRESHOLD else 0)


def _line_chunks_from_tesseract(preprocessed: Image.Image) -> list[tuple[str, float]]:
    """Return (line_text, confidence 0..1 per line mean of non-zero tesseract word confidences)."""
    data = pytesseract.image_to_data(preprocessed, output_type=pytesseract.Output.DICT)
    n = len(data.get("text", []))
    order: list[tuple[int, int, int]] = []
    seen: set[tuple[int, int, int]] = set()
    buckets: dict[tuple[int, int, int], list[tuple[str, float]]] = defaultdict(list)

    for i in range(n):
        token = (data["text"][i] or "").strip()
        if not token:
            continue
        key = (int(data["block_num"][i]), int(data["par_num"][i]), int(data["line_num"][i]))
        if key not in seen:
            seen.add(key)
            order.append(key)
        try:
            cf = float(str(data["conf"][i]))
        except (TypeError, ValueError):
            cf = -1.0
        buckets[key].append((token, cf))

    result: list[tuple[str, float]] = []
    for key in order:
        words = buckets[key]
        line_text = " ".join(w for w, _ in words).strip()
        confs = [c for _, c in words if c > 0]
        avg = (sum(confs) / len(confs) / 100.0) if confs else 0.0
        avg = max(0.0, min(1.0, avg))
        if len(line_text) >= MIN_BLOCK_CHARS:
            result.append((line_text, avg))
    return result


def _split_pdf_paragraphs(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    parts = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    if len(parts) <= 1:
        parts = [p.strip() for p in text.split("\n") if p.strip()]
    return [p for p in parts if len(p) >= MIN_BLOCK_CHARS]


def _ingest_pdf(path: Path, doc_name: str) -> list[TextBlock]:
    entries: list[tuple[str, int, float, str]] = []
    with pdfplumber.open(path) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            try:
                raw = (page.extract_text() or "").strip()
                if len(raw) > PDF_SCANNED_MAX_CHARS:
                    for para in _split_pdf_paragraphs(raw):
                        entries.append((para, page_idx, 1.0, "direct_parse"))
                else:
                    pil_img = page.to_image(resolution=300).original
                    pre = _preprocess_for_ocr(pil_img)
                    for line_text, conf in _line_chunks_from_tesseract(pre):
                        entries.append((line_text, page_idx, conf, "ocr_scan"))
            except Exception as exc:
                logger.warning("PDF page %s failed for %s: %s", page_idx, doc_name, exc)
                continue
    return _to_text_blocks(entries, doc_name)


def _ingest_docx(path: Path, doc_name: str) -> list[TextBlock]:
    entries: list[tuple[str, int, float, str]] = []
    try:
        doc = DocxDocument(path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) >= MIN_BLOCK_CHARS:
                entries.append((text, 0, 1.0, "docx_parse"))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if len(text) >= MIN_BLOCK_CHARS:
                        entries.append((text, 0, 1.0, "docx_parse"))
    except Exception as exc:
        logger.warning("DOCX ingestion failed for %s: %s", doc_name, exc)
    return _to_text_blocks(entries, doc_name)


def _ingest_image_path(path: Path, doc_name: str) -> list[TextBlock]:
    entries: list[tuple[str, int, float, str]] = []
    try:
        with Image.open(path) as img:
            img = img.convert("RGB")
            pre = _preprocess_for_ocr(img)
            for line_text, conf in _line_chunks_from_tesseract(pre):
                entries.append((line_text, 1, conf, "ocr_image"))
    except Exception as exc:
        logger.warning("Image ingestion failed for %s: %s", doc_name, exc)
    return _to_text_blocks(entries, doc_name)


def _to_text_blocks(entries: list[tuple[str, int, float, str]], doc_name: str) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    for text, page_no, conf, method in entries:
        clean = text.strip()
        if len(clean) < MIN_BLOCK_CHARS:
            continue
        blocks.append(
            TextBlock(
                id=f"TB-{uuid.uuid4().hex[:12]}",
                doc_name=doc_name,
                page_number=page_no,
                block_index=len(blocks),
                text=clean,
                ocr_confidence=conf,
                extraction_method=method,
            ),
        )
    return blocks


def ingest_document_sync(file_path: str, doc_name: str) -> list[TextBlock]:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_DOC_EXTS.union(SUPPORTED_IMAGE_EXTS):
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return _ingest_pdf(path, doc_name)
    if ext == ".docx":
        return _ingest_docx(path, doc_name)
    return _ingest_image_path(path, doc_name)


async def ingest_document(file_path: str, doc_name: str) -> list[TextBlock]:
    return await asyncio.to_thread(ingest_document_sync, file_path, doc_name)


async def ingest_upload(upload_file: UploadFile) -> list[TextBlock]:
    upload_root = _upload_root()
    await asyncio.to_thread(upload_root.mkdir, parents=True, exist_ok=True)

    suffix = Path(upload_file.filename or "upload.bin").suffix.lower()
    safe_suffix = suffix if suffix else ".bin"
    dest = upload_root / f"{uuid.uuid4().hex}{safe_suffix}"

    body = await upload_file.read()

    def _write() -> None:
        dest.write_bytes(body)

    await asyncio.to_thread(_write)
    name = upload_file.filename or dest.name

    try:
        return await ingest_document(str(dest), name)
    finally:
        def _unlink() -> None:
            try:
                dest.unlink(missing_ok=True)
            except OSError:
                pass

        await asyncio.to_thread(_unlink)


class IngestionEngine:
    SUPPORTED_EXTENSIONS = SUPPORTED_DOC_EXTS | SUPPORTED_IMAGE_EXTS

    def ingest_file_bytes(self, *, file_name: str, content: bytes) -> list[TextBlock]:
        extension = Path(file_name).suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        suffix = extension if extension else ".bin"
        with tempfile.NamedTemporaryFile(prefix="verdict-ingest-", suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            return ingest_document_sync(tmp_path, file_name)
        finally:
            Path(tmp_path).unlink(missing_ok=True)
